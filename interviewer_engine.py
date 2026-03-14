# -*- coding: utf-8 -*-
"""
睿聘智模 v3.1  ·  interviewer_engine.py
核心引擎：无终端 I/O，所有方法返回 dict/str/bytes，由 app.py 调用。

【评分公式（已修正）】
  1. 每个维度下可能有多道题，先对该维度所有题目的 dimension_score 求平均值
  2. 再将各维度平均分按权重加权求和 → 最终得分（严格在 0-100 区间）
  最终得分 = Σ( avg(dim_i_scores) × weight_i )   其中 Σweight_i = 1.0
"""

import json
import re
import math
from datetime import datetime
from io import BytesIO
from openai import OpenAI

from config import (
    API_BASE, API_KEY, MODEL_NAME, LLM_TIMEOUT,
    INTERVIEWER_NAME, INTERVIEWER_TITLE, INTERVIEWER_YEARS_EXP, INTERVIEWER_STYLE,
    SCORE_GRADES, INTERVIEW_MODES, DIMENSION_WEIGHT_REFERENCE,
    DEFAULT_MAX_FOLLOWUP_ROUNDS, DEFAULT_DEPTH_SCORE_THRESHOLD,
    DEFAULT_MIN_STAR_ELEMENTS, DEFAULT_INTERVIEW_DURATION_MINUTES, DEFAULT_INTERVIEW_MODE,
)

# ================================================================
#  LLM 单例客户端
# ================================================================
_client = None

def get_client() -> OpenAI:
    global _client
    if _client is None:
        _client = OpenAI(api_key=API_KEY, base_url=API_BASE, timeout=LLM_TIMEOUT)
    return _client


def call_llm(messages: list, temperature: float = 0.7, max_tokens: int = 4000) -> str:
    try:
        r = get_client().chat.completions.create(
            model=MODEL_NAME, messages=messages,
            temperature=temperature, max_tokens=max_tokens,
        )
        return r.choices[0].message.content.strip()
    except Exception as e:
        return f"[LLM_ERROR] {e}"


def call_llm_json(messages: list, temperature: float = 0.1) -> dict:
    for _ in range(3):
        try:
            raw = call_llm(messages, temperature=temperature, max_tokens=4096)
            if raw.startswith("[LLM_ERROR]"):
                continue
            cleaned = re.sub(r'```(?:json)?\s*', '', raw)
            cleaned = re.sub(r'```\s*$', '', cleaned).strip()
            s, e = cleaned.find('{'), cleaned.rfind('}') + 1
            if s != -1 and e > s:
                return json.loads(cleaned[s:e])
        except Exception:
            pass
    return {}


# ================================================================
#  图片简历 OCR
# ================================================================
def extract_text_from_image(b64_data: str, mime_type: str) -> str:
    msgs = [{"role": "user", "content": [
        {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{b64_data}"}},
        {"type": "text", "text":
            "请完整提取这张图片中的所有文字（这是一份简历），"
            "保留原有结构与层次，不添加任何额外说明，直接输出纯文本。"}
    ]}]
    return call_llm(msgs, temperature=0.05, max_tokens=3000)


# ================================================================
#  评分工具
# ================================================================
def score_to_grade(score: float) -> tuple:
    for lo, hi, grade, desc in SCORE_GRADES:
        if lo <= score <= hi:
            return grade, desc
    return "D", "不达标"


def compute_weighted_score(evaluations: list) -> tuple[float, dict]:
    """
    正确的加权评分：
      1. 按维度分组，求每维度所有题目的平均分
      2. 平均分 × 权重求和 = 最终加权总分
    返回 (weighted_total, dim_averages_dict)
    """
    dim_buckets: dict[str, dict] = {}
    for ev in evaluations:
        dn = ev.get("dimension_name", "")
        dw = float(ev.get("dimension_weight", 0.2))
        sc = float(ev.get("evaluation", {}).get("dimension_score", 0))
        if dn not in dim_buckets:
            dim_buckets[dn] = {"weight": dw, "scores": []}
        dim_buckets[dn]["scores"].append(sc)

    weighted_total = 0.0
    dim_averages = {}
    for dn, data in dim_buckets.items():
        avg = sum(data["scores"]) / len(data["scores"])
        dim_averages[dn] = round(avg, 1)
        weighted_total += avg * data["weight"]

    return round(min(weighted_total, 100.0), 1), dim_averages


# ================================================================
#  提示词库
# ================================================================

def _mode_style(mode: str) -> str:
    return INTERVIEW_MODES.get(mode, INTERVIEW_MODES["standard"]).get("style_prompt", "")

def _mode_aggression(mode: str) -> str:
    return INTERVIEW_MODES.get(mode, INTERVIEW_MODES["standard"]).get("followup_aggression", "moderate")


# ── System Prompt（人格锚定）──────────────────────────────────────
def build_system_prompt(mode: str, max_followup: int) -> str:
    mode_name  = INTERVIEW_MODES.get(mode, {}).get("name", "标准专业")
    mode_style = _mode_style(mode)
    return f"""你是{INTERVIEWER_NAME}，拥有{INTERVIEWER_YEARS_EXP}年经验的{INTERVIEWER_TITLE}，业内以"{INTERVIEWER_STYLE}"著称。
当前采用【{mode_name}】面试模式。

━━━━━━━━ 核心评估方法论 ━━━━━━━━

① STAR+D 结构化追问框架
   Situation（情境）：背景、规模、复杂度
   Task（任务）：候选人需达到的目标与挑战
   Action（行动）：候选人"个人"具体做了什么（核心：他/她，非"我们"）
   Result（结果）：量化产出与影响
   Depth（深度）："为什么这样做"背后的思维依据

② 冰山模型
   水面以上（显性）：技能/知识/经历 → 简历可初步验证
   水面以下（隐性）：思维方式/动机/价值观/成长潜力 → 仅面试能揭示

③ Reflexion 自我反思（每次追问前必须在心里完成）
   "候选人回答的真实意图是什么？他/她回避了什么？"
   "我的问题是否造成歧义？追问最有价值的切入点是哪个STAR要素？"

━━━━━━━━ 当前模式风格要求 ━━━━━━━━

{mode_style}

━━━━━━━━ 对话铁则 ━━━━━━━━

· 始终第一人称"我"，绝不暴露AI身份、评分机制
· 对候选人每段回答先有一句真实具体的回应，禁止"非常好！""您说得对！"等空洞夸赞
· 追问永远不重复原问题措辞，从回答"信息裂缝"处切入
· 最多追问 {max_followup} 轮即自然推进，不执着于填满STAR模板
· 候选人偏题时温和引导，给完成表达的机会
"""


# ── 岗位识别 + 五维权重（LLM自由确定，参考预设）────────────────
def prompt_identify_position(jd: str) -> str:
    ref_text = "\n".join(
        f"  · {k}：{v}" for k, v in DIMENSION_WEIGHT_REFERENCE.items()
    )
    return f"""你是资深人才评估专家，精通岗位分析与五维能力模型设计。

请深度分析以下岗位JD，完成：
① 识别岗位类型与资历层级
② 为五个能力维度独立确定权重

【岗位JD】
{jd}

【五维维度（固定顺序）】
D1 = 专业技能与知识（水面以上）
D2 = 问题解决与创新思维（水面交界）
D3 = 沟通协作与影响力（水面以下）
D4 = 学习成长与适应力（水面以下）
D5 = 价值观与文化契合度（水面深层）

【各类岗位的历史参考权重（仅供参考，请根据此JD的实际特性独立判断，不必照搬）】
{ref_text}
维度顺序对应：D1 D2 D3 D4 D5

【权重设计要求】
- 五个权重之和必须精确等于 1.00（保留两位小数）
- 必须体现差异化，体现该岗位对各维度的真实依赖（避免均匀分配）
- 综合考虑岗位层级、团队规模、业务属性独立判断
- 你的判断优先于参考值；若此岗位有特殊性，应体现出来

严格按以下JSON格式输出，无任何额外文字：

{{
  "position_type": "技术研发",
  "position_type_reasoning": "一句话说明岗位类型识别依据",
  "seniority_level": "中级",
  "weights": {{
    "D1_专业技能与知识": 0.35,
    "D2_问题解决与创新思维": 0.27,
    "D3_沟通协作与影响力": 0.15,
    "D4_学习成长与适应力": 0.14,
    "D5_价值观与文化契合度": 0.09
  }},
  "weight_reasoning": {{
    "D1": "权重设定理由（1句话）",
    "D2": "理由",
    "D3": "理由",
    "D4": "理由",
    "D5": "理由"
  }},
  "key_competencies": ["此岗位最核心的胜任力1", "胜任力2", "胜任力3"],
  "red_flags_to_probe": ["需重点验证的潜在风险1", "风险2"]
}}
"""


# ── 动态题数分配 ──────────────────────────────────────────────────
def prompt_determine_question_count(weights: dict, duration: int,
                                     position_type: str, seniority: str) -> str:
    mn = max(5, duration // 15)
    mx = min(18, duration // 8)
    return f"""你是面试设计专家。请根据以下信息，为五个维度合理分配题目数量。

【面试时长】{duration} 分钟
【岗位类型】{position_type}  |  【资历层级】{seniority}
【五维权重】{json.dumps(weights, ensure_ascii=False)}
（维度顺序：D1专业技能 D2问题解决 D3沟通协作 D4学习成长 D5价值观）

【分配原则】
1. 总题数控制在 {mn}~{mx} 道（每题含追问平均约 {duration//mx}-{duration//mn} 分钟）
2. 每个维度最少 1 道，最多 4 道
3. 权重越高的维度分配题目越多；权重低的维度象征性覆盖（1题）
4. 五个维度题数之和 = 总题数

严格按JSON格式输出，无任何额外文字：

{{
  "total_questions": 10,
  "allocation": {{"D1": 3, "D2": 2, "D3": 2, "D4": 2, "D5": 1}},
  "reasoning": "一句话说明分配逻辑"
}}
"""


# ── 构建完整面试方案 ──────────────────────────────────────────────
def prompt_build_plan(jd: str, resume: str, weights: dict,
                      position_type: str, seniority: str,
                      key_competencies: list, red_flags: list,
                      q_allocation: dict) -> str:
    wv = list(weights.values())
    w_lines = "\n".join(f"   · {k}：权重 {v:.0%}" for k, v in weights.items())
    alloc_str = " | ".join(f"D{i+1}={q_allocation.get(f'D{i+1}',2)}题" for i in range(5))

    return f"""你是顶级面试策略分析师，请构建精准的结构化面试评估方案。

━━━━ 输入 ━━━━

【岗位JD】
{jd}

【候选人简历】
{resume}

【岗位分析】
类型：{position_type}  层级：{seniority}
核心胜任力：{', '.join(key_competencies)}
需探查风险：{', '.join(red_flags) if red_flags else '暂无'}

【五维权重（已确定）】
{w_lines}

【各维度题目数量】{alloc_str}

━━━━ 输出质量要求（违反均不合格）━━━━

1. 每个维度按上方数量精确输出对应道数的问题
2. 每道问题必须引用候选人简历中真实的公司/项目/技能名称定制，禁止"某项目""某公司"
3. 问题措辞像真人面试官对话，不像问卷条目
4. 简历有疑点（频繁跳槽/时间断层/成果过于完美）时设计具体验证性问题
5. follow_up_hints 指向该题最可能的薄弱处，不写通用方向
6. opening_greeting 自然提及候选人简历中某个具体亮点作为破冰点

严格按JSON格式输出，无任何额外文字：

{{
  "position_title": "岗位名称",
  "match_score": 0.75,
  "match_analysis": "简历与JD核心匹配度分析（引用具体内容，150字以内）",
  "candidate_highlights": ["亮点1（引用简历具体内容）", "亮点2", "亮点3"],
  "candidate_gaps": ["不足/疑点1（具体指出）", "疑点2"],
  "interview_strategy": "整体策略说明（150字以内）",
  "opening_greeting": "以面试官身份的开场白，自然提及简历某个具体亮点，说明今天流程，专业温暖，100-120字",
  "dimensions": [
    {{
      "dimension_id": 1,
      "name": "专业技能与知识",
      "iceberg_level": "水面以上",
      "weight": {wv[0] if len(wv)>0 else 0.28},
      "focus_points": ["具体考察方向1", "方向2"],
      "dimension_strategy": "针对此候选人的评估策略，一句话",
      "questions": [
        {{
          "question_id": "D1Q1",
          "question_text": "结合候选人简历真实经历定制的问题（必须引用具体项目/公司，不要杜撰简历中不存在的项目和实习经历）",
          "knowledge_point": "核心考察知识/能力点",
          "star_focus": "Action+Result",
          "expected_depth": "优秀回答应达到的深度（具体描述）",
          "transition": "",
          "follow_up_hints": ["若回答表浅时的具体追问方向", "若有疑点时的验证性追问"]
        }}
      ]
    }},
    {{
      "dimension_id": 2, "name": "问题解决与创新思维", "iceberg_level": "水面交界",
      "weight": {wv[1] if len(wv)>1 else 0.24},
      "focus_points": [], "dimension_strategy": "", "questions": []
    }},
    {{
      "dimension_id": 3, "name": "沟通协作与影响力", "iceberg_level": "水面以下",
      "weight": {wv[2] if len(wv)>2 else 0.22},
      "focus_points": [], "dimension_strategy": "", "questions": []
    }},
    {{
      "dimension_id": 4, "name": "学习成长与适应力", "iceberg_level": "水面以下",
      "weight": {wv[3] if len(wv)>3 else 0.15},
      "focus_points": [], "dimension_strategy": "", "questions": []
    }},
    {{
      "dimension_id": 5, "name": "价值观与文化契合度", "iceberg_level": "水面深层",
      "weight": {wv[4] if len(wv)>4 else 0.11},
      "focus_points": [], "dimension_strategy": "", "questions": []
    }}
  ]
}}

注意：维度2-5的 questions 数组按上方题数分配填充，结构与维度1相同，必须结合简历定制，不能留空数组。
"""


# ── 评估候选人回答（100分制）────────────────────────────────────
def prompt_evaluate_answer(question: str, answer: str, context: str,
                            dimension_name: str, dim_weight: float,
                            expected_depth: str, followup_round: int,
                            evals_summary: str, depth_threshold: int,
                            max_followup: int, min_star: int) -> str:
    is_last = (followup_round >= max_followup)
    return f"""你是结构化面试评估专家，精通STAR+D方法论与冰山理论。
请对候选人回答进行100分制的深度结构化评估。

━━━━ 评估输入 ━━━━

【评估维度】{dimension_name}（权重 {dim_weight:.0%}）
【面试官提问】{question}
【期望回答深度】{expected_depth}
【候选人本次回答】{answer}
【当前追问轮次】第{followup_round}轮（0=首次回答）| 最后机会：{"是" if is_last else "否"}

【近期对话上下文】
{context}

【已完成维度评估摘要（跨维度参考）】
{evals_summary if evals_summary else "（暂无，这是第一道题）"}

━━━━ 评分标准（100分整数）━━━━

depth_score：
  0-30  仅有结论无过程，或严重偏题
  31-54 有基本描述，缺个人视角与深度
  55-64 结构基本完整，有个人判断，缺量化或细节
  65-79 结构清晰，个人视角明确，有量化，细节充分
  80-100 卓越：独特洞察+逻辑严密+量化结果+反思复盘

dimension_score：综合STAR完整度、回答深度、与维度期望的匹配程度，给出0-100整数

follow_up_needed 判断：
  · STAR要素出现少于{min_star}个 → true
  · depth_score < {depth_threshold} → true
  · 已是第{max_followup}轮 → 强制false
  · 候选人已给出充分深度回答 → false

严格按JSON格式输出，无任何额外文字：

{{
  "star_analysis": {{
    "Situation": {{
      "present": true, "quality_pct": 70,
      "extracted_content": "从回答提取的情境原话或近似原话",
      "note": "质量评估：情境是否清晰？是否有规模/背景数据？"
    }},
    "Task": {{
      "present": true, "quality_pct": 60,
      "extracted_content": "提取内容",
      "note": "候选人是否清晰界定了自己的职责边界和具体挑战？"
    }},
    "Action": {{
      "present": false, "quality_pct": 0,
      "extracted_content": "（未提及或高度模糊）",
      "note": "缺失原因：如'回答停留在团队行为，未说明本人具体做了什么'"
    }},
    "Result": {{
      "present": false, "quality_pct": 0,
      "extracted_content": "（未提及）",
      "note": "是否有量化结果？还是仅停留在主观感受？"
    }}
  }},
  "star_completeness_score": 50,
  "depth_score": 42,
  "answer_level": "superficial",
  "iceberg_analysis": {{
    "surface_signals": ["水面以上能力信号（结合具体内容）"],
    "deep_signals": ["水面以下信号：思维方式/动机/价值观（结合具体措辞）"],
    "cross_dimension_pattern": "与已完成维度对比的一致性规律或矛盾信号",
    "concerns": ["值得警惕的具体信号，如回答高度标准化、与简历描述矛盾等"]
  }},
  "dimension_score": 55,
  "follow_up_needed": true,
  "follow_up_reason": "精确说明追问核心原因（指向STAR哪个要素、什么缺失）",
  "follow_up_focus": "追问最有价值的具体切入点（一句话）",
  "key_insights": ["关键洞察1（结合候选人具体表达）", "洞察2"],
  "raw_assessment": "综合文字评估，100字以内，专业直接"
}}
"""


# ── Reflexion 追问生成 ──────────────────────────────────────────
def prompt_generate_followup(original_q: str, answer: str, evaluation: dict,
                              followup_round: int, context: str, mode: str,
                              max_followup: int) -> str:
    star     = evaluation.get("star_analysis", {})
    missing  = [k for k, v in star.items() if isinstance(v,dict) and not v.get("present",True)]
    weak     = [k for k, v in star.items() if isinstance(v,dict) and v.get("present",True) and v.get("quality_pct",100)<55]
    focus    = evaluation.get("follow_up_focus", "回答的具体行动与量化结果")
    reason   = evaluation.get("follow_up_reason", "缺乏足够的个人行动细节")
    concerns = evaluation.get("iceberg_analysis", {}).get("concerns", [])
    is_last  = (followup_round + 1 >= max_followup)
    aggr     = _mode_aggression(mode)
    aggr_inst = {"low":"语气轻松温和，用开放引导",
                 "moderate":"语气专业自然，像朋友间深度交流",
                 "high":"语气直接，可直接指出回答不足，甚至提出反驳，测试思维清晰度"}.get(aggr,"")

    return f"""你是{INTERVIEWER_NAME}，正在面试追问环节。你已完成Reflexion内部反思，现在输出追问话语。

【内部反思结论（勿输出）】
原始问题：{original_q}
候选人回答摘要：{answer[:500]}
STAR缺失要素：{', '.join(missing) if missing else '无'}
STAR质量偏低要素：{', '.join(weak) if weak else '无'}
追问核心原因：{reason}
最有价值追问切入点：{focus}
可疑信号：{'; '.join(concerns) if concerns else '无'}
第{followup_round+1}轮追问 | 是否是最后机会：{"是" if is_last else "否"}

【近期对话上下文】
{context}

【当前模式追问要求】
{aggr_inst}
{"最后一次追问：核心是确认最关键的疑点，可以更直接" if is_last else ""}

【生成规则】
① 1句真实具体的回应（禁止空洞夸赞）
② 1-2句针对"{focus}"的精准追问，使用引导语切入
③ 整体不超过90字，不要任何标签/换行/前缀

直接输出追问内容：
"""


# ── 自然过渡到下一个问题 ────────────────────────────────────────
def prompt_transition_question(next_q_info: dict, last_answer: str,
                                context: str, done_q: int, total_q: int,
                                mode: str) -> str:
    return f"""你是{INTERVIEWER_NAME}，刚结束一道题，需要自然地过渡到下一个问题。

【候选人最后一段回答】
{last_answer[:400]}

【近期对话摘要】
{context}

【面试进度】已完成{done_q}题 / 共{total_q}题

【下一个问题】{next_q_info.get('question_text','')}
【预设过渡方向（仅参考）】{next_q_info.get('transition','')}

【当前模式风格】
{_mode_style(mode)}

请生成"过渡+提问"话语：
① 1-2句自然过渡（从候选人刚才某个词/观点作为桥梁，或维度切换引导）
②直接提出下一个问题（保持原意，可调整措辞使其更自然） 
整体不超过130字，不要出现"现在进入下一个环节"等机械语，也不要展现出来你的内心活动

直接输出话语（无标签）：
"""


# ── 本题分析报告 ─────────────────────────────────────────────────
def prompt_question_report(question: str, answer: str, evaluation: dict,
                            dimension_name: str, dim_weight: float,
                            followup_hist: list, evals_summary: str) -> str:
    fh = ""
    for i, h in enumerate(followup_hist, 1):
        fh += f"\n  第{i}轮追问：{h.get('question','')}\n  候选人回应：{h.get('answer','')[:300]}\n"

    return f"""你是资深面试评估专家，生成精准的本题分析报告（面试官内部参考）。

【维度】{dimension_name}（权重{dim_weight:.0%}） | 【问题】{question}
【主要回答】{answer[:600]}
【追问记录】{fh if fh else '（无追问）'}
【量化评估】STAR完整度：{evaluation.get('star_completeness_score',0)}/100 | 深度：{evaluation.get('depth_score',0)}/100 | 维度分：{evaluation.get('dimension_score',0)}/100
【冰山分析】{json.dumps(evaluation.get('iceberg_analysis',{}), ensure_ascii=False)}
【关键洞察】{json.dumps(evaluation.get('key_insights',[]), ensure_ascii=False)}
【已完成维度参考】{evals_summary if evals_summary else '（暂无）'}

请生成四部分报告（用【】标题，总计240-320字）。
所有结论必须引用候选人具体回答内容（原话或近似原话）作为佐证。
如有多轮追问，分析主回答与追问回答的差异及候选人在压力下的表现变化。
如与已完成维度存在规律性信号需在报告中点出。

【STAR要素评估】各要素存在/质量/亮点/缺口，引用原话。
【能力表现信号】观察到的具体能力表现（水面以上）。
【冰山洞察】水面以下：思维模式/动机/价值观/矛盾信号（结合具体表达）。
【本题小结与决策参考】综合判断1-2句，说明对录用决策的参考价值方向。
"""


# ── 最终综合评估报告（含公司培养路线）──────────────────────────
def prompt_final_report(jd: str, resume: str, plan: dict,
                        position_type: str, weights: dict,
                        evaluations: list, dialogue_log: str,
                        interview_mode_name: str,
                        weighted_total: float, dim_averages: dict) -> str:
    w_lines = "\n".join(f"   · {k}：{v:.0%}" for k, v in weights.items())
    grade, grade_desc = score_to_grade(weighted_total)

    dim_detail = []
    for dn, avg in dim_averages.items():
        g, _ = score_to_grade(avg)
        dw = next((ev.get("dimension_weight",0) for ev in evaluations if ev.get("dimension_name")==dn), 0)
        q_cnt = sum(1 for ev in evaluations if ev.get("dimension_name")==dn)
        insights = []
        for ev in evaluations:
            if ev.get("dimension_name") == dn:
                insights.extend(ev.get("evaluation",{}).get("key_insights",[])[:2])
        dim_detail.append({
            "维度": dn, "权重": f"{dw:.0%}", "题数": q_cnt,
            "平均分": f"{avg:.1f}/100", "等级": g,
            "关键洞察": insights[:3]
        })

    return f"""你是资深HR总监{INTERVIEWER_NAME}，为关键岗位录用决策撰写最终面试评估报告（{interview_mode_name}模式）。
报告将直接影响录用决策，并作为公司制定培养方案的核心依据。

━━━━ 核心评估数据 ━━━━

【岗位】{plan.get('position_title','目标岗位')} | 【类型】{position_type}
【简历-JD匹配度】{plan.get('match_score',0):.0%}
【五维权重（LLM动态生成）】
{w_lines}

【各维度评估结果（评分公式：先算各维度平均分，再加权求和）】
{json.dumps(dim_detail, ensure_ascii=False, indent=2)}

【加权综合得分】{weighted_total:.1f} / 100.0 → {grade}级（{grade_desc}）
  计算逻辑：Σ(各维度平均分 × 维度权重)，权重之和=1，总分严格在0-100区间

【面试对话记录节选】
{dialogue_log[:5000]}

━━━━ 报告结构（七个部分，每部分用【 】标题）━━━━

所有结论必须基于具体面试行为/表达作为佐证，禁止通用化表述。

【一、综合能力评分卡】
逐维度列出（表格形式）：
  维度名称 | 权重 | 题数 | 各题得分 | 维度平均分 | 等级 | 核心判断
最后：
  加权综合得分：{weighted_total:.1f}/100
  综合等级：{grade}级 — {grade_desc}
  等级说明：A级(85-100)卓越 | B级(70-84)优秀 | C级(55-69)合格 | D级(0-54)不达标

【二、候选人完整能力画像】
· 水面以上（显性层）：专业能力质量、技术/业务深度、工作经验成色的综合评价
· 水面以下（隐性层）：从面试行为、应对压力方式、措辞选择中识别的思维特质、动机结构、价值取向
· 整体人才标签（2-3个关键词）

【三、核心优势（3-4条）】
格式：▸ [优势名称]：具体描述 + 面试原话或具体行为佐证 + 对岗位的实际价值

【四、潜在风险与能力缺口（2-3条）】
格式：▸ [风险名称]：具体描述 + 面试观察信号 + 入职后可能影响 + 管控建议

【五、录用决策建议】
▸ 决策结论（从以下四个中选一）：
  ✅ 强烈推荐录用 / ✓ 推荐录用 / ◎ 建议进行补充面试 / ✗ 不建议录用
▸ 核心决策依据（3-5条，每条一句话，有具体面试行为支撑）
▸ 特别注意事项或录用条件（如有）
▸ 若推荐录用：薪资谈判策略 + 试用期核心观察方向
▸ 若建议补充面试：明确指出需补充考察的1-2个具体维度

【六、新员工入职融合支持方案】
▸ 前30天关键支持：基于候选人特质，哪些需主动引导/避免哪些误区
▸ 工作任务切入建议：适合从哪类任务开始，哪类需搭档支持
▸ 团队融合注意点：基于协作风格，如何帮助与现有团队建立信任

【七、公司人才培养路线（6-12个月）】
以公司和直属领导视角制定：

  ▸ 第一阶段（0-3个月 · 定向补强期）
  — 公司需重点投入补强的能力短板（结合候选人具体缺口）
  — 建议配置的资源：培训方向/导师类型/刻意练习场景
  — 公司可安排的学习型任务
  — 阶段里程碑（可观察/可验证）

  ▸ 第二阶段（3-6个月 · 实战提升期）
  — 公司应提供的成长机会和挑战性任务（具体项目类型/角色机会）
  — 核心能力提升目标（具体量化或可观察）
  — 评估节点：如何判断候选人是否达到正式胜任水平

  ▸ 第三阶段（6-12个月 · 潜力释放期）
  — 公司下一步可提供的发展路径（晋升/横向扩展/专家方向）
  — 公司需提前储备的资源或机会
  — 12个月整体评估标准：从"达标"到"卓越"的分级判断依据

报告结尾注明：面试日期（{datetime.now().strftime('%Y年%m月%d日')}）| 岗位：{plan.get('position_title','')} | 主面试官：{INTERVIEWER_NAME}
"""


# ================================================================
#  InterviewerEngine 核心引擎
# ================================================================

class InterviewerEngine:
    """无终端 I/O，所有状态服务端维护，每个 session 独立一个实例。"""

    def __init__(self, settings: dict | None = None):
        s = settings or {}
        # 面试参数（可覆盖默认值）
        self.max_followup       = int(s.get("max_followup_rounds", DEFAULT_MAX_FOLLOWUP_ROUNDS))
        self.depth_threshold    = int(s.get("depth_score_threshold", DEFAULT_DEPTH_SCORE_THRESHOLD))
        self.min_star           = int(s.get("min_star_elements", DEFAULT_MIN_STAR_ELEMENTS))
        self.duration_minutes   = int(s.get("duration_minutes", DEFAULT_INTERVIEW_DURATION_MINUTES))
        self.interview_mode     = s.get("interview_mode", DEFAULT_INTERVIEW_MODE)

        # 输入
        self.jd      = ""
        self.resume  = ""
        # 分析结果
        self.plan             = {}
        self.position_type    = "通用"
        self.seniority        = "中级"
        self.weights          = {}
        self.key_competencies = []
        self.red_flags        = []
        self.q_allocation     = {}
        # 面试状态
        self.chat_hist        = []
        self.dial_log         = []
        self.evaluations      = []
        self.all_questions    = []   # [(dim_idx, q_info), ...]
        self.current_q_idx    = 0
        self.followup_count   = 0
        self.followup_history = []
        self.last_answer      = ""
        self.last_followup_text = ""
        # 报告
        self.final_report_text = ""
        self.system_prompt     = build_system_prompt(self.interview_mode, self.max_followup)

    # ── Step 1: 岗位分析 ──────────────────────────────────────
    def analyze_position(self) -> dict:
        msgs = [
            {"role": "system", "content": "你是专业人才评估专家，请严格按JSON格式输出，无任何额外文字。"},
            {"role": "user",   "content": prompt_identify_position(self.jd)},
        ]
        r = call_llm_json(msgs, temperature=0.1)
        if r and "weights" in r:
            self.position_type    = r.get("position_type", "通用")
            self.seniority        = r.get("seniority_level", "中级")
            self.weights          = r.get("weights", {})
            self.key_competencies = r.get("key_competencies", [])
            self.red_flags        = r.get("red_flags_to_probe", [])
            reasoning             = r.get("weight_reasoning", {})
        else:
            self.position_type    = "通用"
            self.seniority        = "中级"
            ref = DIMENSION_WEIGHT_REFERENCE.get("通用", [0.28,0.24,0.22,0.15,0.11])
            keys = ["D1_专业技能与知识","D2_问题解决与创新思维","D3_沟通协作与影响力","D4_学习成长与适应力","D5_价值观与文化契合度"]
            self.weights = dict(zip(keys, ref))
            self.key_competencies = []
            self.red_flags        = []
            reasoning             = {}

        # 归一化权重
        total = sum(self.weights.values())
        if total > 0 and abs(total - 1.0) > 0.005:
            self.weights = {k: round(v/total, 4) for k, v in self.weights.items()}

        return {"position_type": self.position_type, "seniority": self.seniority,
                "weights": self.weights, "key_competencies": self.key_competencies,
                "red_flags": self.red_flags, "weight_reasoning": reasoning}

    # ── Step 2: 题数分配 ──────────────────────────────────────
    def determine_question_count(self) -> dict:
        msgs = [
            {"role": "system", "content": "你是面试设计专家，请严格按JSON格式输出，无任何额外文字。"},
            {"role": "user",   "content": prompt_determine_question_count(
                self.weights, self.duration_minutes, self.position_type, self.seniority)},
        ]
        r = call_llm_json(msgs, temperature=0.1)
        if r and "allocation" in r:
            self.q_allocation = r.get("allocation", {})
            total = r.get("total_questions", 10)
        else:
            self.q_allocation = {f"D{i+1}": 2 for i in range(5)}
            total = 10
        return {"allocation": self.q_allocation, "total": total,
                "reasoning": r.get("reasoning", "") if r else ""}

    # ── Step 3: 构建面试方案 ──────────────────────────────────
    def build_plan(self) -> dict:
        msgs = [
            {"role": "system", "content": "你是顶级面试策略分析师，请严格按JSON格式输出，无任何额外文字。"},
            {"role": "user",   "content": prompt_build_plan(
                self.jd, self.resume, self.weights,
                self.position_type, self.seniority,
                self.key_competencies, self.red_flags, self.q_allocation)},
        ]
        self.plan = call_llm_json(msgs, temperature=0.2)
        if not self.plan or "dimensions" not in self.plan:
            self.plan = self._default_plan()

        # 同步动态权重
        wv = list(self.weights.values())
        for i, dim in enumerate(self.plan.get("dimensions", [])):
            if i < len(wv):
                dim["weight"] = wv[i]

        # 展平题目列表
        self.all_questions = []
        for di, dim in enumerate(self.plan.get("dimensions", [])):
            for qi in dim.get("questions", []):
                self.all_questions.append((di, qi))

        self.chat_hist = [{"role": "system", "content": self.system_prompt}]
        return self.plan

    # ── Step 4: 获取当前问题 ─────────────────────────────────
    def get_current_question_display(self) -> dict:
        idx = self.current_q_idx
        if idx >= len(self.all_questions):
            return {"done": True, "text": ""}

        di, qi = self.all_questions[idx]
        dim = self.plan["dimensions"][di]
        q_text = qi.get("question_text", "")

        if idx == 0:
            opening = self.plan.get("opening_greeting",
                                    f"您好！我是{INTERVIEWER_NAME}，今天由我来主持面试，我们开始吧？")
            display = opening + "\n\n" + q_text
        else:
            msgs = [
                {"role": "system", "content": self.system_prompt},
                {"role": "user",   "content": prompt_transition_question(
                    qi, self.last_answer, self._context_summary(8),
                    idx, len(self.all_questions), self.interview_mode)},
            ]
            display = call_llm(msgs, temperature=0.75, max_tokens=350)
            if not display or display.startswith("[LLM_ERROR]"):
                t = qi.get("transition", "")
                display = (t + " " + q_text).strip() if t else q_text

        self._log("interviewer", display, "question")
        return {
            "done": False, "text": display,
            "question_id": qi.get("question_id", ""),
            "dimension": dim.get("name", ""),
            "dim_weight": dim.get("weight", 0.2),
            "progress": {"current": idx + 1, "total": len(self.all_questions)},
        }

    # ── Step 5: 处理主回答 ───────────────────────────────────
    def process_answer(self, answer: str) -> dict:
        self._log("candidate", answer, "answer")
        self.last_answer = answer
        di, qi = self.all_questions[self.current_q_idx]
        dim = self.plan["dimensions"][di]

        ev = self._evaluate(qi.get("question_text",""), answer, dim, qi, 0)
        self.last_evaluation = ev
        q_report = self._q_report(qi.get("question_text",""), answer, ev,
                                   dim.get("name",""), dim.get("weight",0.2))
        if self._need_followup(ev):
            fu = self._do_followup(qi.get("question_text",""), answer, ev)
            return {"action":"followup","followup_text":fu,
                    "evaluation":self._safe_summary(ev),"question_report":q_report}

        self._archive(di, dim, qi, answer, ev, q_report)
        self.current_q_idx  += 1
        self.followup_count  = 0
        self.followup_history = []
        action = "complete" if self.current_q_idx >= len(self.all_questions) else "next_question"
        return {"action":action,"evaluation":self._safe_summary(ev),"question_report":q_report}

    # ── Step 5b: 处理追问回答 ────────────────────────────────
    def process_followup_answer(self, answer: str) -> dict:
        self._log("candidate", answer, "followup_answer")
        di, qi = self.all_questions[self.current_q_idx]
        dim = self.plan["dimensions"][di]
        self.followup_history.append({"question": self.last_followup_text, "answer": answer})
        self.followup_count += 1

        combined = self.last_answer + f"\n[第{self.followup_count}轮追问补充]: " + answer
        ev = self._evaluate(qi.get("question_text",""), combined, dim, qi, self.followup_count)
        self.last_evaluation = ev
        q_report = self._q_report(qi.get("question_text",""), combined, ev,
                                   dim.get("name",""), dim.get("weight",0.2))
        if self._need_followup(ev):
            fu = self._do_followup(qi.get("question_text",""), combined, ev)
            return {"action":"followup","followup_text":fu,
                    "evaluation":self._safe_summary(ev),"question_report":q_report}

        self._archive(di, dim, qi, combined, ev, q_report)
        self.current_q_idx  += 1
        self.followup_count  = 0
        self.followup_history = []
        action = "complete" if self.current_q_idx >= len(self.all_questions) else "next_question"
        return {"action":action,"evaluation":self._safe_summary(ev),"question_report":q_report}

    # ── Step 6: 生成最终报告 ─────────────────────────────────
    def generate_final_report(self) -> dict:
        # 结束语
        closing_msgs = [
            {"role": "system", "content": self.system_prompt},
            {"role": "user",   "content":
                f"请以{INTERVIEWER_NAME}身份生成面试结束语："
                "真诚感谢候选人（可具体提及某个印象深刻的分享）→简述后续流程→给予具体鼓励。"
                "不要透露任何评估结论。90字以内。"},
        ]
        closing = call_llm(closing_msgs, temperature=0.75, max_tokens=250)
        if not closing or closing.startswith("[LLM_ERROR]"):
            closing = (f"非常感谢您今天与我交流，分享了很多有价值的经历和思考。"
                       "后续HR团队会在3-5个工作日内与您联系，请保持手机畅通。祝您一切顺利！")
        self._log("interviewer", closing, "closing")

        # 正确评分
        weighted_total, dim_averages = compute_weighted_score(self.evaluations)
        grade, grade_desc = score_to_grade(weighted_total)

        # 最终报告文本
        mode_name = INTERVIEW_MODES.get(self.interview_mode, {}).get("name", "标准专业")
        msgs = [
            {"role": "system", "content":
                f"你是资深HR总监{INTERVIEWER_NAME}，正在撰写专业评估报告，"
                "所有结论必须有具体面试行为佐证，培养路线以公司视角制定。"},
            {"role": "user", "content": prompt_final_report(
                self.jd, self.resume, self.plan,
                self.position_type, self.weights, self.evaluations,
                self._format_dial_log(), mode_name,
                weighted_total, dim_averages)},
        ]
        report_text = call_llm(msgs, temperature=0.3, max_tokens=5000)
        self.final_report_text = report_text

        return {
            "closing":        closing,
            "report":         report_text,
            "weighted_score": weighted_total,
            "dim_averages":   dim_averages,
            "grade":          grade,
            "grade_desc":     grade_desc,
            "evaluations":    self.evaluations,
        }

    # ── 导出：完整 HTML ────────────────────────────────────────
    def export_html(self) -> str:
        weighted_total, dim_averages = compute_weighted_score(self.evaluations)
        grade, grade_desc = score_to_grade(weighted_total)
        mode_name = INTERVIEW_MODES.get(self.interview_mode, {}).get("name", "标准专业")
        now_str   = datetime.now().strftime("%Y年%m月%d日 %H:%M")

        # 维度得分卡 HTML
        dim_cards_html = ""
        from collections import defaultdict
        dim_qs = defaultdict(list)
        for ev in self.evaluations:
            dim_qs[ev.get("dimension_name","")].append(ev)

        for dn, evs in dim_qs.items():
            avg = dim_averages.get(dn, 0)
            dg, _ = score_to_grade(avg)
            dw = evs[0].get("dimension_weight", 0) if evs else 0
            dg_color = {"A":"#27ae60","B":"#2980b9","C":"#f39c12","D":"#e74c3c"}.get(dg,"#95a5a6")

            q_rows_html = ""
            for i, ev in enumerate(evs, 1):
                e   = ev.get("evaluation", {})
                sc  = e.get("dimension_score", 0)
                star = e.get("star_analysis", {})
                star_html = ""
                for elem in ["Situation","Task","Action","Result"]:
                    sd = star.get(elem, {})
                    q  = sd.get("quality_pct", 0)
                    bar_col = ("#27ae60" if q>=65 else ("#f39c12" if q>=40 else "#e74c3c"))
                    star_html += (f'<span style="margin-right:10px;font-size:12px;color:{bar_col}">'
                                  f'{"✓" if sd.get("present") else "✗"} {elem} {q}%</span>')
                insights = "; ".join(e.get("key_insights",[])[:2])
                q_rows_html += f"""
                <tr>
                  <td style="padding:8px;border-bottom:1px solid #f5f5f5;font-size:12px;color:#666">Q{i}</td>
                  <td style="padding:8px;border-bottom:1px solid #f5f5f5;font-size:12px">{ev.get('question','')[:60]}…</td>
                  <td style="padding:8px;border-bottom:1px solid #f5f5f5;text-align:center;font-size:16px;font-weight:700;color:#1a73e8">{sc}</td>
                  <td style="padding:8px;border-bottom:1px solid #f5f5f5;font-size:12px">{star_html}</td>
                  <td style="padding:8px;border-bottom:1px solid #f5f5f5;font-size:12px;color:#666">{insights}</td>
                </tr>"""

            report_text = evs[-1].get("question_report","") if evs else ""
            dim_cards_html += f"""
            <div style="background:#fff;border-radius:10px;border:1px solid #e9ecef;margin-bottom:16px;overflow:hidden">
              <div style="background:#f8f9fa;padding:14px 20px;display:flex;align-items:center;gap:16px;border-bottom:1px solid #e9ecef">
                <div style="font-weight:700;font-size:15px;flex:1">{dn}</div>
                <div style="font-size:12px;color:#666">权重 {dw:.0%} · {len(evs)}题</div>
                <div style="font-size:28px;font-weight:800;color:#1a73e8">{avg:.1f}</div>
                <div style="font-size:12px;color:#999">/100</div>
                <div style="background:{dg_color};color:#fff;padding:3px 10px;border-radius:12px;font-size:12px;font-weight:700">{dg}级</div>
              </div>
              <div style="overflow-x:auto">
                <table style="width:100%;border-collapse:collapse">
                  <thead><tr style="background:#f8f9fa">
                    <th style="padding:8px;text-align:left;font-size:12px;color:#666">题号</th>
                    <th style="padding:8px;text-align:left;font-size:12px;color:#666">问题</th>
                    <th style="padding:8px;text-align:center;font-size:12px;color:#666">得分</th>
                    <th style="padding:8px;text-align:left;font-size:12px;color:#666">STAR要素</th>
                    <th style="padding:8px;text-align:left;font-size:12px;color:#666">关键洞察</th>
                  </tr></thead>
                  <tbody>{q_rows_html}</tbody>
                </table>
              </div>
              {f'<div style="padding:14px 20px;background:#fffde7;font-size:13px;line-height:1.8;color:#444;border-top:1px solid #f0f2f5;white-space:pre-wrap">{report_text}</div>' if report_text else ''}
            </div>"""

        # 对话记录 HTML
        chat_html = ""
        for entry in self.dial_log:
            role    = entry.get("role","")
            content = entry.get("content","").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace("\n","<br>")
            if role == "interviewer":
                chat_html += f'<div style="margin-bottom:14px"><div style="font-size:11px;color:#999;margin-bottom:4px">🎙 {INTERVIEWER_NAME}</div><div style="background:#e8f0fe;padding:12px 16px;border-radius:12px 12px 12px 4px;font-size:13px;line-height:1.7;max-width:85%">{content}</div></div>'
            else:
                chat_html += f'<div style="margin-bottom:14px;display:flex;flex-direction:column;align-items:flex-end"><div style="font-size:11px;color:#999;margin-bottom:4px">👤 候选人</div><div style="background:#f8f9fa;padding:12px 16px;border-radius:12px 12px 4px 12px;font-size:13px;line-height:1.7;max-width:85%">{content}</div></div>'

        # 最终报告 HTML（格式化渲染）
        report_html = ""
        for line in self.final_report_text.split("\n"):
            s = line.strip()
            if s.startswith("【") and "】" in s:
                report_html += f'<h3 style="font-size:16px;font-weight:700;color:#1a73e8;margin:20px 0 8px;padding-bottom:6px;border-bottom:2px solid #e8f0fe">{s}</h3>'
            elif s.startswith("▸"):
                report_html += f'<p style="margin:8px 0;padding-left:16px;border-left:3px solid #1a73e8;color:#333;line-height:1.8">{s}</p>'
            elif s.startswith("—"):
                report_html += f'<p style="margin:4px 0 4px 24px;color:#555;font-size:13px;line-height:1.8">{s}</p>'
            elif "强烈推荐录用" in s:
                report_html += f'<p style="background:#e8f5e9;color:#27ae60;padding:10px 16px;border-radius:8px;font-weight:700;margin:8px 0">{s}</p>'
            elif "推荐录用" in s and "不建议" not in s:
                report_html += f'<p style="background:#e3f2fd;color:#1565c0;padding:10px 16px;border-radius:8px;font-weight:700;margin:8px 0">{s}</p>'
            elif "不建议录用" in s:
                report_html += f'<p style="background:#fce8e6;color:#c62828;padding:10px 16px;border-radius:8px;font-weight:700;margin:8px 0">{s}</p>'
            elif "补充面试" in s:
                report_html += f'<p style="background:#fff8e1;color:#e65100;padding:10px 16px;border-radius:8px;font-weight:700;margin:8px 0">{s}</p>'
            elif s:
                report_html += f'<p style="margin:6px 0;color:#333;line-height:1.8">{s}</p>'

        score_color = {"A":"#27ae60","B":"#2980b9","C":"#f39c12","D":"#e74c3c"}.get(grade,"#666")

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>面试评估报告 · {self.plan.get('position_title','岗位')} · {now_str}</title>
<style>
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'PingFang SC','Microsoft YaHei',sans-serif;background:#f0f2f5;color:#202124;font-size:14px;line-height:1.6}}
  @media print{{body{{background:#fff}}.no-print{{display:none}}.page-break{{page-break-before:always}}}}
</style>
</head>
<body>
<div style="max-width:1000px;margin:0 auto;padding:24px">

  <!-- 封面 -->
  <div style="background:linear-gradient(135deg,#1a73e8,#0d47a1);color:#fff;border-radius:16px;padding:48px;margin-bottom:24px;text-align:center">
    <div style="font-size:13px;opacity:.7;margin-bottom:8px">睿聘智模 · 面试综合评估报告</div>
    <h1 style="font-size:28px;font-weight:800;margin-bottom:8px">{self.plan.get('position_title','目标岗位')}</h1>
    <div style="font-size:15px;opacity:.85;margin-bottom:24px">{self.position_type} · {self.seniority} · {mode_name}模式</div>
    <div style="display:inline-block;background:rgba(255,255,255,.15);border-radius:20px;padding:24px 60px">
      <div style="font-size:72px;font-weight:800;line-height:1;color:{score_color if score_color!='#666' else '#fff'}">{weighted_total:.1f}</div>
      <div style="font-size:18px;opacity:.8;margin-top:4px">/ 100 分</div>
      <div style="font-size:22px;font-weight:700;margin-top:8px">{grade}级 — {grade_desc}</div>
    </div>
    <div style="font-size:12px;opacity:.6;margin-top:20px">面试官：{INTERVIEWER_NAME} · {now_str}</div>
  </div>

  <!-- 权重分配 -->
  <div style="background:#fff;border-radius:12px;box-shadow:0 2px 8px rgba(0,0,0,.08);margin-bottom:20px;overflow:hidden">
    <div style="background:#1a73e8;color:#fff;padding:14px 24px;font-size:15px;font-weight:700">⚖️ 五维权重分配（LLM动态生成）</div>
    <div style="padding:20px;display:flex;flex-wrap:wrap;gap:12px">
      {"".join(f'<div style="flex:1;min-width:150px;background:#f8f9fa;border-radius:8px;padding:14px"><div style="font-size:12px;color:#666;margin-bottom:4px">{k}</div><div style="font-size:24px;font-weight:700;color:#1a73e8">{v:.0%}</div><div style="height:6px;background:#e9ecef;border-radius:3px;margin-top:6px;overflow:hidden"><div style="height:100%;width:{v*100:.0f}%;background:#1a73e8;border-radius:3px"></div></div></div>' for k,v in self.weights.items())}
    </div>
  </div>

  <!-- 各维度评估 -->
  <div style="background:#fff;border-radius:12px;box-shadow:0 2px 8px rgba(0,0,0,.08);margin-bottom:20px;overflow:hidden">
    <div style="background:#1a73e8;color:#fff;padding:14px 24px;font-size:15px;font-weight:700">📊 各维度面试评估（评分 = 各题平均 × 权重）</div>
    <div style="padding:20px">{dim_cards_html}</div>
  </div>

  <!-- 对话记录 -->
  <div style="background:#fff;border-radius:12px;box-shadow:0 2px 8px rgba(0,0,0,.08);margin-bottom:20px;overflow:hidden">
    <div style="background:#1a73e8;color:#fff;padding:14px 24px;font-size:15px;font-weight:700">💬 完整面试对话记录</div>
    <div style="padding:24px">{chat_html}</div>
  </div>

  <!-- 最终报告 -->
  <div style="background:#fff;border-radius:12px;box-shadow:0 2px 8px rgba(0,0,0,.08);margin-bottom:20px;overflow:hidden">
    <div style="background:#1a73e8;color:#fff;padding:14px 24px;font-size:15px;font-weight:700">📋 最终综合评估报告</div>
    <div style="padding:28px">{report_html}</div>
  </div>

  <div style="text-align:center;color:#999;font-size:12px;padding:20px">
    © 睿聘智模 v3.1 · {now_str} · 主面试官：{INTERVIEWER_NAME}
  </div>
</div>
</body>
</html>"""

    # ── 导出：PDF（reportlab 完整多页）─────────────────────────
    def export_pdf(self) -> bytes:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib.colors import HexColor, white, black
            from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                            Table, TableStyle, PageBreak, HRFlowable,
                                            KeepTogether)
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
        except ImportError:
            return b""

        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=2*cm, rightMargin=2*cm,
            topMargin=2.5*cm, bottomMargin=2*cm,
            title=f"面试评估报告_{self.plan.get('position_title','')}",
            author=INTERVIEWER_NAME,
        )

        # ── 颜色定义 ──
        C_BLUE  = HexColor("#1a73e8")
        C_LBLUE = HexColor("#e8f0fe")
        C_GREEN = HexColor("#27ae60")
        C_LGRN  = HexColor("#e8f5e9")
        C_ORG   = HexColor("#f39c12")
        C_RED   = HexColor("#e74c3c")
        C_GRAY  = HexColor("#5f6368")
        C_LGRAY = HexColor("#f8f9fa")
        C_DARK  = HexColor("#202124")

        # ── 样式 ──
        ss = getSampleStyleSheet()
        def S(name, **kw):
            return ParagraphStyle(name, parent=ss["Normal"], **kw)

        sTitle   = S("T", fontSize=24, fontName="Helvetica-Bold", textColor=white,
                     alignment=TA_CENTER, spaceAfter=4)
        sSubtitle= S("S", fontSize=13, fontName="Helvetica", textColor=white,
                     alignment=TA_CENTER, spaceAfter=12)
        sMeta    = S("M", fontSize=10, fontName="Helvetica", textColor=HexColor("#ccddff"),
                     alignment=TA_CENTER)
        sSection = S("Sec", fontSize=14, fontName="Helvetica-Bold", textColor=C_BLUE,
                     spaceBefore=16, spaceAfter=6)
        sSub     = S("Sub", fontSize=12, fontName="Helvetica-Bold", textColor=C_DARK,
                     spaceBefore=10, spaceAfter=4)
        sBody    = S("B", fontSize=10, fontName="Helvetica", textColor=C_DARK,
                     leading=16, spaceAfter=4)
        sBodyB   = S("BB", fontSize=10, fontName="Helvetica-Bold", textColor=C_DARK,
                     leading=16, spaceAfter=4)
        sBullet  = S("Blt", fontSize=10, fontName="Helvetica", textColor=C_DARK,
                     leftIndent=12, leading=16, spaceAfter=3)
        sCaption = S("Cap", fontSize=9, fontName="Helvetica", textColor=C_GRAY,
                     alignment=TA_CENTER, spaceAfter=4)
        sGreen   = S("Grn", fontSize=11, fontName="Helvetica-Bold", textColor=C_GREEN,
                     backColor=C_LGRN, leading=16, spaceAfter=4)
        sRed     = S("Red", fontSize=11, fontName="Helvetica-Bold", textColor=C_RED,
                     leading=16, spaceAfter=4)

        weighted_total, dim_averages = compute_weighted_score(self.evaluations)
        grade, grade_desc = score_to_grade(weighted_total)
        mode_name = INTERVIEW_MODES.get(self.interview_mode, {}).get("name", "标准专业")
        now_str   = datetime.now().strftime("%Y年%m月%d日 %H:%M")

        story = []

        # ═══════════════════════════════════
        # 封面页
        # ═══════════════════════════════════
        story.append(Spacer(1, 3*cm))
        cover_data = [
            [Paragraph("睿聘智模  ·  面试综合评估报告", S("ct",fontSize=11,textColor=C_GRAY,alignment=TA_CENTER))],
            [Paragraph(self.plan.get("position_title","目标岗位"), S("ct",fontSize=22,fontName="Helvetica-Bold",textColor=C_BLUE,alignment=TA_CENTER,spaceAfter=8))],
            [Paragraph(f"{self.position_type}  ·  {self.seniority}  ·  {mode_name}模式",
                       S("ct",fontSize=12,textColor=C_GRAY,alignment=TA_CENTER))],
            [Spacer(1, 0.5*cm)],
            [Table([[Paragraph(f"{weighted_total:.1f}", S("sc",fontSize=52,fontName="Helvetica-Bold",
                              textColor=C_BLUE,alignment=TA_CENTER))],
                    [Paragraph("/ 100  综合得分", S("sc2",fontSize=12,textColor=C_GRAY,alignment=TA_CENTER))],
                    [Paragraph(f"{grade}级  —  {grade_desc}",
                               S("sc3",fontSize=14,fontName="Helvetica-Bold",
                                 textColor=C_GREEN if grade in["A","B"] else C_RED,
                                 alignment=TA_CENTER))],
                   ],
                   colWidths=[12*cm],
                   style=TableStyle([
                       ("ALIGN",  (0,0),(0,-1),"CENTER"),
                       ("VALIGN", (0,0),(0,-1),"MIDDLE"),
                       ("BOX",    (0,0),(-1,-1),2,C_BLUE),
                       ("ROUNDEDCORNERS",[8]),
                       ("BACKGROUND",(0,0),(-1,-1),C_LBLUE),
                       ("TOPPADDING",(0,0),(-1,-1),12),
                       ("BOTTOMPADDING",(0,0),(-1,-1),12),
                   ]))],
            [Spacer(1,0.5*cm)],
            [Paragraph(f"面试官：{INTERVIEWER_NAME}  ·  {now_str}",
                       S("ct",fontSize=10,textColor=C_GRAY,alignment=TA_CENTER))],
        ]
        cover_table = Table(cover_data, colWidths=[15*cm])
        cover_table.setStyle(TableStyle([("ALIGN",(0,0),(-1,-1),"CENTER"),
                                          ("VALIGN",(0,0),(-1,-1),"MIDDLE")]))
        story.append(cover_table)
        story.append(PageBreak())

        # ═══════════════════════════════════
        # 一、简历匹配分析
        # ═══════════════════════════════════
        story.append(Paragraph("一、简历 × JD 匹配分析", sSection))
        story.append(HRFlowable(width="100%", thickness=1, color=C_BLUE, spaceAfter=8))
        pct = int(self.plan.get("match_score", 0.7) * 100)
        story.append(Paragraph(f"匹配度：{pct}%", sBodyB))
        story.append(Paragraph(self.plan.get("match_analysis", ""), sBody))
        story.append(Spacer(1, 0.2*cm))

        if self.plan.get("candidate_highlights"):
            story.append(Paragraph("候选人亮点：", sSub))
            for h in self.plan.get("candidate_highlights", []):
                story.append(Paragraph(f"✦  {h}", sBullet))

        if self.plan.get("candidate_gaps"):
            story.append(Paragraph("待验证疑点：", sSub))
            for g in self.plan.get("candidate_gaps", []):
                story.append(Paragraph(f"△  {g}", sBullet))

        # ═══════════════════════════════════
        # 二、五维权重
        # ═══════════════════════════════════
        story.append(Spacer(1, 0.4*cm))
        story.append(Paragraph("二、五维能力权重分配（LLM动态生成）", sSection))
        story.append(HRFlowable(width="100%", thickness=1, color=C_BLUE, spaceAfter=8))

        w_data = [["维度", "权重", "题目数", "权重理由"]]
        for i,(k,v) in enumerate(self.weights.items()):
            di_key = f"D{i+1}"
            q_cnt  = self.q_allocation.get(di_key, "—")
            w_data.append([k, f"{v:.0%}", str(q_cnt), ""])
        w_table = Table(w_data, colWidths=[6*cm,2*cm,2*cm,5*cm])
        w_table.setStyle(TableStyle([
            ("BACKGROUND", (0,0),(-1,0), C_BLUE),
            ("TEXTCOLOR",  (0,0),(-1,0), white),
            ("FONTNAME",   (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",   (0,0),(-1,-1), 9),
            ("ALIGN",      (1,0),(2,-1), "CENTER"),
            ("GRID",       (0,0),(-1,-1), 0.5, HexColor("#e9ecef")),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[white, C_LGRAY]),
            ("TOPPADDING", (0,0),(-1,-1), 5),
            ("BOTTOMPADDING",(0,0),(-1,-1), 5),
        ]))
        story.append(w_table)

        # ═══════════════════════════════════
        # 三、各维度评估得分卡（正确公式说明）
        # ═══════════════════════════════════
        story.append(Spacer(1, 0.4*cm))
        story.append(Paragraph("三、各维度评估得分（评分公式：先求各维度平均分，再加权求和）", sSection))
        story.append(HRFlowable(width="100%", thickness=1, color=C_BLUE, spaceAfter=6))
        story.append(Paragraph(
            "计算逻辑：每个维度可能有多道题，先对该维度所有题目的得分求平均值，再乘以权重并求和，"
            "确保总分严格在 0-100 区间内。", sBody))

        # 汇总表
        score_data = [["维度", "权重", "题数", "各题得分", "维度平均分", "等级", "加权贡献"]]
        from collections import defaultdict
        dim_qs_map = defaultdict(list)
        for ev in self.evaluations:
            dim_qs_map[ev.get("dimension_name","")].append(ev)

        for dn, evs in dim_qs_map.items():
            avg  = dim_averages.get(dn, 0)
            dg, _= score_to_grade(avg)
            dw   = evs[0].get("dimension_weight", 0) if evs else 0
            scores_str = ", ".join(str(int(ev.get("evaluation",{}).get("dimension_score",0))) for ev in evs)
            contrib = avg * dw
            score_data.append([
                dn[:10], f"{dw:.0%}", str(len(evs)), scores_str,
                f"{avg:.1f}", f"{dg}级", f"{contrib:.2f}",
            ])
        score_data.append(["合计", "100%", "", "", "", "", f"{weighted_total:.1f}"])

        sc_table = Table(score_data, colWidths=[4.2*cm,1.4*cm,1.2*cm,2.5*cm,2.2*cm,1.2*cm,2*cm])
        sc_table.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(-1,0),  C_BLUE),
            ("TEXTCOLOR",     (0,0),(-1,0),  white),
            ("FONTNAME",      (0,0),(-1,0),  "Helvetica-Bold"),
            ("FONTNAME",      (0,-1),(-1,-1),"Helvetica-Bold"),
            ("BACKGROUND",    (0,-1),(-1,-1),C_LBLUE),
            ("FONTSIZE",      (0,0),(-1,-1),  8),
            ("ALIGN",         (1,0),(-1,-1),  "CENTER"),
            ("GRID",          (0,0),(-1,-1),  0.5, HexColor("#e9ecef")),
            ("ROWBACKGROUNDS",(0,1),(-1,-2), [white, C_LGRAY]),
            ("TOPPADDING",    (0,0),(-1,-1),  5),
            ("BOTTOMPADDING", (0,0),(-1,-1),  5),
        ]))
        story.append(sc_table)

        # ═══════════════════════════════════
        # 四、各维度详细评估
        # ═══════════════════════════════════
        story.append(PageBreak())
        story.append(Paragraph("四、各维度详细面试评估", sSection))
        story.append(HRFlowable(width="100%", thickness=1, color=C_BLUE, spaceAfter=8))

        for dn, evs in dim_qs_map.items():
            avg  = dim_averages.get(dn, 0)
            dg, _= score_to_grade(avg)
            dw   = evs[0].get("dimension_weight",0) if evs else 0
            block = []
            block.append(Paragraph(f"▌ {dn}  权重{dw:.0%}  |  维度平均分：{avg:.1f}/100  {dg}级", sSub))

            for i, ev in enumerate(evs, 1):
                e   = ev.get("evaluation", {})
                sc  = e.get("dimension_score", 0)
                block.append(Paragraph(f"  Q{i}：{ev.get('question','')[:80]}", sBodyB))
                block.append(Paragraph(f"  本题得分：{sc}/100  |  深度：{e.get('depth_score',0)}/100  |  STAR完整度：{e.get('star_completeness_score',0)}/100", sBody))
                star = e.get("star_analysis",{})
                for elem in ["Situation","Task","Action","Result"]:
                    sd = star.get(elem,{})
                    icon = "✓" if sd.get("present") else "✗"
                    block.append(Paragraph(f"  {icon} {elem}：{sd.get('quality_pct',0)}%  {sd.get('note','')[:60]}", sBullet))
                insights = e.get("key_insights",[])
                if insights:
                    block.append(Paragraph(f"  关键洞察：{'; '.join(insights[:2])}", sBody))
                concerns = e.get("iceberg_analysis",{}).get("concerns",[])
                if concerns:
                    block.append(Paragraph(f"  ⚑ 注意信号：{'; '.join(concerns[:2])}", sBullet))

            if evs:
                q_report = evs[-1].get("question_report","")
                if q_report:
                    block.append(Paragraph("  本维度分析报告：", sBodyB))
                    for line in q_report.split("\n"):
                        s = line.strip()
                        if s:
                            block.append(Paragraph(f"  {s}", sBody))

            story.append(KeepTogether(block[:8]))
            for item in block[8:]:
                story.append(item)
            story.append(Spacer(1, 0.3*cm))

        # ═══════════════════════════════════
        # 五、完整对话记录
        # ═══════════════════════════════════
        story.append(PageBreak())
        story.append(Paragraph("五、完整面试对话记录", sSection))
        story.append(HRFlowable(width="100%", thickness=1, color=C_BLUE, spaceAfter=8))

        for entry in self.dial_log:
            role    = entry.get("role","")
            content = entry.get("content","")[:600]
            if role == "interviewer":
                story.append(Paragraph(f"🎙 {INTERVIEWER_NAME}：", sBodyB))
                story.append(Paragraph(content, sBody))
            else:
                story.append(Paragraph("👤 候选人：", sBodyB))
                story.append(Paragraph(content, sBullet))
            story.append(Spacer(1, 0.15*cm))

        # ═══════════════════════════════════
        # 六、最终综合评估报告正文
        # ═══════════════════════════════════
        story.append(PageBreak())
        story.append(Paragraph("六、最终综合评估报告", sSection))
        story.append(HRFlowable(width="100%", thickness=1, color=C_BLUE, spaceAfter=8))

        for line in self.final_report_text.split("\n"):
            s = line.strip()
            if not s:
                story.append(Spacer(1, 0.1*cm))
            elif s.startswith("【") and "】" in s:
                story.append(Spacer(1,0.2*cm))
                story.append(Paragraph(s, sSub))
                story.append(HRFlowable(width="100%",thickness=0.5,color=C_LBLUE,spaceAfter=4))
            elif s.startswith("▸"):
                story.append(Paragraph(s, sBullet))
            elif s.startswith("—"):
                story.append(Paragraph(f"    {s}", sBullet))
            elif "强烈推荐录用" in s:
                story.append(Paragraph(f"✅  {s}", sGreen))
            elif "不建议录用" in s:
                story.append(Paragraph(f"✗  {s}", sRed))
            else:
                story.append(Paragraph(s, sBody))

        # 页脚
        story.append(Spacer(1, 0.5*cm))
        story.append(HRFlowable(width="100%", thickness=0.5, color=C_GRAY))
        story.append(Paragraph(
            f"睿聘智模 v3.1  ·  面试官：{INTERVIEWER_NAME}  ·  {now_str}",
            S("ft",fontSize=9,textColor=C_GRAY,alignment=TA_CENTER)))

        doc.build(story)
        return buf.getvalue()

    # ── 导出：DOCX ─────────────────────────────────────────────
    def export_docx(self) -> bytes:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            doc.core_properties.author = INTERVIEWER_NAME

            weighted_total, dim_averages = compute_weighted_score(self.evaluations)
            grade, grade_desc = score_to_grade(weighted_total)
            mode_name = INTERVIEW_MODES.get(self.interview_mode,{}).get("name","标准专业")

            def h(txt, lvl=1):
                doc.add_heading(txt, level=lvl)

            def p(txt, bold=False):
                para = doc.add_paragraph()
                run  = para.add_run(txt)
                run.bold = bold

            h(f"面试综合评估报告  ·  {self.plan.get('position_title','')}", 1)
            p(f"面试官：{INTERVIEWER_NAME}  |  模式：{mode_name}  |  {datetime.now().strftime('%Y年%m月%d日')}")
            p(f"加权综合得分：{weighted_total:.1f}/100.0  |  综合等级：{grade}级 — {grade_desc}", bold=True)
            doc.add_paragraph()

            h("一、简历 × JD 匹配分析", 2)
            p(f"匹配度：{int(self.plan.get('match_score',0.7)*100)}%", bold=True)
            p(self.plan.get("match_analysis",""))
            for item in self.plan.get("candidate_highlights",[]):
                doc.add_paragraph(f"✦ {item}", style="List Bullet")
            for item in self.plan.get("candidate_gaps",[]):
                doc.add_paragraph(f"△ {item}", style="List Bullet")

            h("二、五维权重", 2)
            for k,v in self.weights.items():
                p(f"  · {k}：{v:.0%}")

            h("三、各维度评估得分（正确公式）", 2)
            p("计算逻辑：先求各维度平均分，再加权求和。最终得分严格在0-100区间。")
            from collections import defaultdict
            dim_qs_map = defaultdict(list)
            for ev in self.evaluations:
                dim_qs_map[ev.get("dimension_name","")].append(ev)
            for dn, evs in dim_qs_map.items():
                avg = dim_averages.get(dn,0)
                dg, _ = score_to_grade(avg)
                dw = evs[0].get("dimension_weight",0) if evs else 0
                scores_str = ", ".join(str(int(ev.get("evaluation",{}).get("dimension_score",0))) for ev in evs)
                p(f"{dn}（权重{dw:.0%}）：各题={scores_str} → 平均={avg:.1f}/100  {dg}级", bold=True)
            p(f"加权综合得分：{weighted_total:.1f}/100.0", bold=True)

            h("四、完整面试对话记录", 2)
            for entry in self.dial_log:
                role    = entry.get("role","")
                content = entry.get("content","")
                label   = INTERVIEWER_NAME if role=="interviewer" else "候选人"
                para = doc.add_paragraph()
                para.add_run(f"【{label}】").bold = True
                para.add_run(f" {content}")

            h("五、最终综合评估报告", 2)
            doc.add_paragraph(self.final_report_text)

            buf = BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except ImportError:
            return b""

    # ── 内部辅助方法 ────────────────────────────────────────────
    def _evaluate(self, question: str, answer: str, dim: dict,
                  qi: dict, followup_round: int) -> dict:
        msgs = [
            {"role":"system","content":"你是专业面试评估专家，精通STAR+D和冰山理论。请严格按JSON格式输出，无任何额外文字。"},
            {"role":"user","content":prompt_evaluate_answer(
                question, answer, self._context_summary(10),
                dim.get("name",""), dim.get("weight",0.2),
                qi.get("expected_depth","结构清晰、有具体行动细节、有量化结果"),
                followup_round, self._evals_summary(),
                self.depth_threshold, self.max_followup, self.min_star)},
        ]
        ev = call_llm_json(msgs, temperature=0.1)
        return ev if ev else self._default_eval()

    def _need_followup(self, ev: dict) -> bool:
        if self.followup_count >= self.max_followup:
            return False
        star = ev.get("star_analysis", {})
        present_n = sum(1 for v in star.values() if isinstance(v,dict) and v.get("present",False))
        if present_n < self.min_star:
            return True
        if ev.get("depth_score", 100) < self.depth_threshold:
            return True
        if not ev.get("follow_up_needed", True):
            return False
        return False

    def _do_followup(self, original_q: str, answer: str, ev: dict) -> str:
        msgs = [
            {"role":"system","content":self.system_prompt},
            {"role":"user","content":prompt_generate_followup(
                original_q, answer, ev, self.followup_count,
                self._context_summary(8), self.interview_mode, self.max_followup)},
        ]
        fu = call_llm(msgs, temperature=0.78, max_tokens=300)
        if not fu or fu.startswith("[LLM_ERROR]"):
            fu = "能否更具体地描述一下您在其中采取的核心行动步骤，以及最终带来了什么可量化的结果？"
        self.last_followup_text = fu
        self._log("interviewer", fu, "followup")
        return fu

    def _q_report(self, question: str, answer: str, ev: dict,
                  dim_name: str, dim_weight: float) -> str:
        msgs = [
            {"role":"system","content":"你是专业人才评估专家，生成精准深度的面试本题分析报告（内部参考）。"},
            {"role":"user","content":prompt_question_report(
                question, answer, ev, dim_name, dim_weight,
                self.followup_history, self._evals_summary())},
        ]
        report = call_llm(msgs, temperature=0.35, max_tokens=900)
        return report if report and not report.startswith("[LLM_ERROR]") else "（报告生成失败）"

    def _archive(self, di: int, dim: dict, qi: dict,
                 answer: str, ev: dict, q_report: str):
        self.evaluations.append({
            "dimension_name":   dim.get("name",""),
            "dimension_id":     dim.get("dimension_id", di+1),
            "dimension_weight": dim.get("weight", 0.2),
            "question":         qi.get("question_text",""),
            "answer":           answer,
            "evaluation":       ev,
            "question_report":  q_report,
            "followup_count":   self.followup_count,
        })

    def _log(self, role: str, content: str, etype: str = ""):
        self.dial_log.append({"role": role, "content": content, "type": etype})
        if role == "interviewer":
            self.chat_hist.append({"role":"assistant","content":content})
        else:
            self.chat_hist.append({"role":"user","content":content})

    def _context_summary(self, n: int = 8) -> str:
        recent = self.dial_log[-n:]
        return "\n".join(
            f"{'面试官' if e['role']=='interviewer' else '候选人'}: "
            f"{e['content'][:280]}{'…' if len(e['content'])>280 else ''}"
            for e in recent
        )

    def _evals_summary(self) -> str:
        if not self.evaluations:
            return ""
        lines = []
        for ev in self.evaluations:
            e  = ev.get("evaluation",{})
            cross = e.get("iceberg_analysis",{}).get("cross_dimension_pattern","")
            lines.append(
                f"· {ev.get('dimension_name','')}（权重{ev.get('dimension_weight',0):.0%}）："
                f"得分{e.get('dimension_score',0):.0f}/100 | 深度{e.get('depth_score',0):.0f} | "
                f"洞察：{'; '.join(e.get('key_insights',[])[:2])}"
                + (f" | 跨维度：{cross}" if cross else "")
            )
        return "\n".join(lines)

    def _format_dial_log(self) -> str:
        return "\n".join(
            f"{'【面试官】' if e['role']=='interviewer' else '【候选人】'}\n{e['content']}\n"
            for e in self.dial_log
        )

    def _safe_summary(self, ev: dict) -> dict:
        star = ev.get("star_analysis",{})
        return {
            "star_summary": {k: {"present":v.get("present",False),
                                  "quality_pct":v.get("quality_pct",0),
                                  "note":v.get("note","")[:60]}
                              for k,v in star.items() if isinstance(v,dict)},
            "star_completeness": ev.get("star_completeness_score",0),
            "depth_score":       ev.get("depth_score",0),
            "dimension_score":   ev.get("dimension_score",0),
            "answer_level":      ev.get("answer_level",""),
            "key_insights":      ev.get("key_insights",[]),
            "concerns":          ev.get("iceberg_analysis",{}).get("concerns",[]),
            "raw_assessment":    ev.get("raw_assessment",""),
        }

    def _default_eval(self) -> dict:
        return {
            "star_analysis": {k: {"present":k in["Situation","Task"],
                                   "quality_pct":50,"extracted_content":"","note":""}
                               for k in ["Situation","Task","Action","Result"]},
            "star_completeness_score": 50, "depth_score": 50, "answer_level": "moderate",
            "iceberg_analysis": {"surface_signals":[],"deep_signals":[],
                                  "cross_dimension_pattern":"","concerns":[]},
            "dimension_score":   60, "follow_up_needed": True,
            "follow_up_reason":  "需要更多具体行动细节",
            "follow_up_focus":   "个人具体行动与量化结果",
            "key_insights":      [], "raw_assessment": "回答基本完整，期待更多具体细节",
        }

    def _default_plan(self) -> dict:
        wv = list(self.weights.values()) if self.weights else [0.28,0.24,0.22,0.15,0.11]
        dims = [
            ("专业技能与知识","水面以上",wv[0] if len(wv)>0 else 0.28,
             "能跟我介绍一个您做过的最有挑战性的项目吗？您在其中承担了什么核心职责，是如何推动它落地的？"),
            ("问题解决与创新思维","水面交界",wv[1] if len(wv)>1 else 0.24,
             "能分享一个您主动发现并解决了一个原本没人意识到的问题的案例吗？"),
            ("沟通协作与影响力","水面以下",wv[2] if len(wv)>2 else 0.22,
             "请分享一次您在没有正式权力的情况下，需要推动一件重要事情的经历。"),
            ("学习成长与适应力","水面以下",wv[3] if len(wv)>3 else 0.15,
             "在职业生涯里，有没有一次您认为自己犯了比较重大错误的经历？那次经历后来对您产生了什么实质性的改变？"),
            ("价值观与文化契合度","水面深层",wv[4] if len(wv)>4 else 0.11,
             "在什么样的工作状态或环境下您会感到最投入、最有能量？能结合具体经历说说吗？"),
        ]
        return {
            "position_title":"目标岗位","match_score":0.70,
            "match_analysis":"简历与JD存在一定匹配度，需通过面试深入验证核心能力。",
            "candidate_highlights":["具备相关工作经验","学历符合岗位要求"],
            "candidate_gaps":["核心技能深度有待验证","项目成果数据有待核实"],
            "interview_strategy":"重点考察专业深度与实际解决问题的能力。",
            "opening_greeting":(f"您好！我是{INTERVIEWER_NAME}，今天由我来主持这次面试，"
                                "不用紧张，把它当作一次深度交流就好，我们开始吧？"),
            "dimensions":[{
                "dimension_id":i+1,"name":n,"iceberg_level":lv,"weight":w,
                "focus_points":[],"dimension_strategy":"",
                "questions":[{
                    "question_id":f"D{i+1}Q1","question_text":q,
                    "knowledge_point":"核心能力","star_focus":"Action+Result",
                    "expected_depth":"结构清晰、有具体行动细节、有量化结果",
                    "transition":"" if i==0 else "换个角度，",
                    "follow_up_hints":["具体行动步骤","量化结果"]
                }]
            } for i,(n,lv,w,q) in enumerate(dims)]
        }